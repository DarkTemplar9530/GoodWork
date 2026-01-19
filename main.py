import streamlit as st
from docx import Document
from docx.shared import Mm, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re

# --- КОНСТАНТЫ ОФОРМЛЕНИЯ ---
FONT_NAME = 'Times New Roman'
FONT_SIZE_MAIN = Pt(14)
FONT_SIZE_TABLE = Pt(12) # Допускается 8-12, берем читаемый
INDENT_MAIN = Cm(1.25)
INDENT_NONE = Cm(0)

def set_page_settings(doc):
    """1. Настройка полей страницы"""
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        # Настройка номера страницы (чтобы он был, но пока пустой)
        section.footer.is_linked_to_previous = False

def clean_formatting(paragraph):
    """Сброс форматирования параграфа перед применением нового"""
    paragraph.paragraph_format.first_line_indent = INDENT_MAIN
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.right_indent = 0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def apply_font_style(paragraph, bold=False, caps=False, size=FONT_SIZE_MAIN, italic=False):
    """Применение шрифта Times New Roman ко всем run'ам в параграфе"""
    # Если нужно сделать CAPS, меняем текст
    if caps:
        text = paragraph.text.upper()
        paragraph.clear()
        paragraph.add_run(text)

    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.bold = bold
        run.font.italic = italic

def set_single_spacing_if_multiline(paragraph):
    """
    10. Названия, содержащие несколько строк, записываются через одинарный интервал.
    Эвристика: если длина текста > 80 символов, считаем, что он займет > 1 строки.
    """
    if len(paragraph.text) > 80:
        paragraph.paragraph_format.line_spacing = 1.0

def add_page_number(doc):
    """
    9. Нумерация страниц (низ, центр).
    Вставка поля {PAGE} через XML.
    """
    for section in doc.sections:
        footer = section.footer
        # Очищаем футер
        for p in footer.paragraphs:
            p.clear()
        
        # Если параграфов нет, создаем
        if not footer.paragraphs:
            footer.add_paragraph()
            
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN

def process_document(uploaded_file):
    doc = Document(uploaded_file)
    set_page_settings(doc)

    # СПИСКИ КЛЮЧЕВЫХ СЛОВ ДЛЯ СТРУКТУРЫ
    STRUCTURAL_HEADERS = [
        "СОДЕРЖАНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", 
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЕ", "ПРИЛОЖЕНИЯ"
    ]

    # Регулярки
    # Глава: "1. НАЗВАНИЕ" или "1 НАЗВАНИЕ" (строго прописные)
    regex_chapter = re.compile(r'^\d+\.?\s+[А-ЯA-Z\s\-\"]+$')
    # Подраздел: "1.1 Название" или "1.1.1 Название"
    regex_subsection = re.compile(r'^\d+(\.\d+)+\s+')
    # Рисунок: "Рисунок 1 – ..."
    regex_figure = re.compile(r'^Рисунок\s+\d+', re.IGNORECASE)
    # Таблица: "Таблица 1 – ..."
    regex_table_caption = re.compile(r'^Таблица\s+\d+', re.IGNORECASE)

    prev_para_was_header = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        clean_formatting(para)
        
        # --- 1. СТРУКТУРНЫЕ ЧАСТИ (СОДЕРЖАНИЕ, ВВЕДЕНИЕ...) ---
        is_struct = False
        for key in STRUCTURAL_HEADERS:
            if key in text.upper() and len(text) < 50: # Проверка длины, чтобы не спутать с текстом
                is_struct = True
                break
        
        if is_struct:
            # Требования: Прописные, Полужирный, По центру, Новая страница
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = INDENT_NONE
            
            # Проверяем, не первая ли это страница (чтобы не рвать титульник)
            if i > 5: # Эвристика: структурные части идут не в начале файла (титульника)
                para.paragraph_format.page_break_before = True
            
            apply_font_style(para, bold=True, caps=True)
            prev_para_was_header = True
            continue

        # --- 2. ГЛАВЫ ОСНОВНОЙ ЧАСТИ (1. НАЗВАНИЕ) ---
        if regex_chapter.match(text):
            # Требования: Прописные, Полужирный, По центру, Новая страница
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = INDENT_NONE
            para.paragraph_format.page_break_before = True
            apply_font_style(para, bold=True, caps=True)
            set_single_spacing_if_multiline(para)
            prev_para_was_header = True
            continue

        # --- 3. ПОДРАЗДЕЛЫ (1.1 Название) ---
        if regex_subsection.match(text):
            # Требования: Строчные (кроме первой), Обычная жирность (НЕ bold), Абзацный отступ
            # Выравнивание по ширине (стандарт текста)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = INDENT_MAIN
            
            # Внимание: Требование п.7 - "Обычной жирности"
            apply_font_style(para, bold=False, caps=False) 
            
            # Отбивка пустой строкой сверху (если предыдущий не заголовок)
            if not prev_para_was_header:
                 para.paragraph_format.space_before = Pt(14) # Примерно одна строка

            set_single_spacing_if_multiline(para)
            prev_para_was_header = True
            continue

        # --- 4. ПОДПИСИ К РИСУНКАМ ---
        if regex_figure.match(text):
            # Требования: По центру, Без отступа, Одинарный интервал, Без точки в конце
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = INDENT_NONE
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(14) # Отбивка от рисунка
            para.paragraph_format.space_after = Pt(14)  # Отбивка от текста
            
            # Убираем точку в конце
            if text.endswith('.'):
                para.text = text[:-1]
                
            apply_font_style(para, bold=False)
            prev_para_was_header = False
            continue

        # --- 5. НАЗВАНИЯ ТАБЛИЦ ---
        if regex_table_caption.match(text):
            # Требования: Слева, Без отступа, Одинарный интервал, Над таблицей
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = INDENT_NONE
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6) # Чуть меньше до самой таблицы
            
            if text.endswith('.'):
                para.text = text[:-1]

            apply_font_style(para, bold=False)
            prev_para_was_header = False
            continue

        # --- 6. ОСНОВНОЙ ТЕКСТ ---
        # Требования: Times New Roman, 14, 1.5 интервал, Выравнивание по ширине, Отступ 1.25
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_font_style(para, bold=False)
        prev_para_was_header = False

    # --- ОБРАБОТКА ТАБЛИЦ ---
    for table in doc.tables:
        table.autofit = False 
        # Можно попытаться растянуть по ширине, но python-docx ограничен в этом.
        
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # В таблицах: 10-12 пт, одинарный интервал, без красной строки
                    p.paragraph_format.first_line_indent = INDENT_NONE
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    
                    for run in p.runs:
                        run.font.name = FONT_NAME
                        run.font.size = FONT_SIZE_TABLE
                        # Сохраняем жирность, если она была (например, в шапке)
                        if run.font.bold:
                            run.font.bold = True

    add_page_number(doc)
    return doc

# --- ИНТЕРФЕЙС STREAMLIT ---
st.set_page_config(page_title="Авто-Нормоконтроль ВКР", layout="centered")

st.title("📄 Автоматическое оформление ВКР/Курсовой")
st.markdown("""
**Сервис приводит документ к стандартам оформления (ГОСТ/УрГУПС):**
1.  **Поля:** 30/10/20/20 мм.
2.  **Шрифт:** Times New Roman, 14 пт, интервал 1.5.
3.  **Заголовки:** ВВЕДЕНИЕ и Главы (1.) — жирные, по центру, с новой страницы.
4.  **Подразделы:** (1.1.) — обычный шрифт, по ширине.
5.  **Таблицы и Рисунки:** Выравнивание подписей, удаление отступов.
6.  **Нумерация страниц:** Внизу по центру.
""")

uploaded_file = st.file_uploader("Загрузите файл .docx (Word)", type="docx")

if uploaded_file is not None:
    if st.button("🚀 Привести в порядок"):
        with st.spinner("Анализирую структуру документа..."):
            try:
                # Обработка
                processed_doc = process_document(uploaded_file)
                
                # Сохранение
                bio = io.BytesIO()
                processed_doc.save(bio)
                bio.seek(0)
                
                st.success("Файл успешно обработан!")
                st.markdown("⚠️ **Важно:** Обязательно проверьте файл после скачивания. Программа не может исправить сложные формулы Microsoft Equation и разрывы внутри таблиц.")
                
                st.download_button(
                    label="📥 Скачать готовый файл",
                    data=bio,
                    file_name=f"fixed_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Ошибка при чтении файла: {e}")
                st.write("Попробуйте сохранить ваш файл как 'Документ Word 2007 (*.docx)' и загрузить снова.")
